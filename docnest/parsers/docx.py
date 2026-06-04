"""
DOCX parser using python-docx.

Extracts headings (Heading 1-6, Title style), paragraphs, and tables from
Word documents in document order — tables appear inside the section they
belong to, not at the end.

python-docx gives explicit heading levels (style names like "Heading 1") which
is more reliable for DOCX files than depth heuristics.

Phase: 1  |  Spec: docs/SPEC_DOCNEST_PYPI.md — Section 10
python-docx docs: https://python-docx.readthedocs.io/
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional, Iterator

from docnest.parsers.base import IParser
from docnest.models import RawDocument, Section, TableData
from docnest.exceptions import ParseError


class DocxParser(IParser):
    """Parses .docx Word documents using python-docx.

    Features:
    - Heading levels from Word styles (Heading 1 → level 1, etc.)
    - Tables extracted in document order with header row preserved
    - Document title from core properties, first Title style, or filename
    - Handles nested headings and multi-level sections

    Usage:
        parser = DocxParser()
        raw = parser.parse("/path/to/doc.docx")
        # raw.sections[n].title  → heading text
        # raw.sections[n].level  → 1-6
        # raw.sections[n].tables → list[TableData]
    """

    # ------------------------------------------------------------------ #
    #  IParser interface                                                   #
    # ------------------------------------------------------------------ #

    def supports(self, file_path: str) -> bool:
        """Return True for .docx files (.doc is not supported by python-docx)."""
        return file_path.lower().endswith(".docx")

    def parse(self, file_path: str) -> RawDocument:
        """Parse a Word document into a RawDocument.

        Args:
            file_path: Absolute or relative path to the .docx file.

        Returns:
            RawDocument with sections and tables extracted.
            Section ids (§N) are NOT assigned — the Normaliser does that.

        Raises:
            ParseError: File missing, corrupted, or python-docx fails.
        """
        path = Path(file_path).resolve()
        if not path.exists():
            raise ParseError(f"DOCX not found: {path}")
        if path.stat().st_size == 0:
            raise ParseError(f"DOCX is empty: {path}")

        try:
            from docx import Document as DocxDocument  # type: ignore[import]
        except ImportError as exc:
            raise ParseError(
                "python-docx is not installed. Run: pip install python-docx"
            ) from exc

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise ParseError(
                f"python-docx failed to open '{path.name}': {exc}"
            ) from exc

        title = self._extract_title(doc, path)
        sections = self._build_sections(doc)

        return RawDocument(
            doc_id=self._make_doc_id(file_path),
            title=title,
            source=str(path),
            format="docx",
            sections=sections,
        )

    # ------------------------------------------------------------------ #
    #  Internal helpers                                                    #
    # ------------------------------------------------------------------ #

    def _extract_title(self, doc: object, path: Path) -> str:
        """Return title from core properties, first Title style, or filename."""
        # 1. Core properties (File → Properties in Word)
        try:
            props_title = doc.core_properties.title.strip()  # type: ignore[attr-defined]
            if props_title:
                return props_title
        except Exception:
            pass

        # 2. First paragraph with 'Title' or 'Heading 1' style
        for para in doc.paragraphs:  # type: ignore[attr-defined]
            style_name = para.style.name
            if style_name in ("Title", "Heading 1") and para.text.strip():
                return para.text.strip()

        # 3. Fall back to filename
        return _filename_to_title(path.stem)

    def _build_sections(self, doc: object) -> list[Section]:
        """Walk paragraphs and tables in document order, grouping by headings.

        python-docx's doc.paragraphs and doc.tables are separate lists.
        To preserve document order we walk the raw XML body directly so that
        tables appear in the section they belong to, not all at the end.

        Rules:
        - Heading 1-6 / Title style  → new Section (formal heading)
        - Pseudo-heading (ALL CAPS / bold short line / colon label) → new Section
        - Normal / Body Text / List  → appended to current Section text
        - Table                      → TableData added to current Section
        - Empty paragraphs           → skipped
        - Content before first heading → implicit 'Introduction' Section (level 1)
        """
        sections: list[Section] = []
        current: Optional[Section] = None
        table_counter = 0

        for block in self._iter_blocks_in_order(doc):
            if _is_paragraph(block):
                style_name = block.style.name  # type: ignore[attr-defined]
                text = block.text.strip()  # type: ignore[attr-defined]
                heading_level = _heading_level(style_name)

                if heading_level is not None and text:
                    # Formal Word Heading style — start a new section
                    if current is not None:
                        current.text = current.text.strip()
                        sections.append(current)
                    current = Section(
                        id="",
                        title=text,
                        level=heading_level,
                        text="",
                    )
                elif text and _is_pseudo_heading(block):  # type: ignore[arg-type]
                    # No formal heading style but looks like a heading
                    # (ALL CAPS, bold short line, or colon-terminated label)
                    if current is not None:
                        current.text = current.text.strip()
                        sections.append(current)
                    current = Section(
                        id="",
                        title=text.rstrip(":").strip(),
                        level=1,
                        text="",
                    )
                else:
                    if not text:
                        continue
                    if current is None:
                        current = Section(
                            id="", title="Introduction", level=1, text=""
                        )
                    if "List" in style_name:
                        current.text += f"- {text}\n"
                    else:
                        current.text += text + "\n\n"

            elif _is_table(block):
                table_counter += 1
                table_data = self._extract_table(block, table_counter)
                if table_data:
                    if current is None:
                        current = Section(id="", title="Tables", level=1, text="")
                    current.tables.append(table_data)

        if current is not None:
            current.text = current.text.strip()
            sections.append(current)

        return sections

    def _iter_blocks_in_order(self, doc: object) -> Iterator[object]:
        """Yield paragraphs and tables in the order they appear in the document.

        Walks the XML body directly using the 'w' namespace tags so that
        tables appear between paragraphs rather than all at the end.
        """
        try:
            from docx.oxml.ns import qn  # type: ignore[import]
            from docx.text.paragraph import Paragraph  # type: ignore[import]
            from docx.table import Table  # type: ignore[import]
        except ImportError:
            return

        body = doc.element.body  # type: ignore[attr-defined]
        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                yield Paragraph(child, doc)  # type: ignore[arg-type]
            elif tag == qn("w:tbl"):
                yield Table(child, doc)  # type: ignore[arg-type]

    def _extract_table(self, table: object, counter: int) -> Optional[TableData]:
        """Convert a python-docx Table to a TableData model.

        The first row is treated as the header row.

        Merged cells: python-docx already returns a full rectangular grid where a
        merged cell repeats its value across every grid position it covers — gridSpan
        (horizontal) repeats across columns, vMerge (vertical) repeats down rows. We
        therefore keep the grid as-is (column alignment preserved, like HTML colspan).
        We must NOT deduplicate: that both misaligned merged columns and collapsed
        legitimately-repeated values (e.g. two cells both "10").
        """
        try:
            rows_raw: list[list[str]] = [
                [cell.text.strip() for cell in row.cells]  # type: ignore[attr-defined]
                for row in table.rows  # type: ignore[attr-defined]
            ]

            if not rows_raw:
                return None

            # Normalise all rows to the same width
            max_cols = max(len(r) for r in rows_raw)
            for r in rows_raw:
                while len(r) < max_cols:
                    r.append("")

            headers = rows_raw[0]
            data_rows = rows_raw[1:]

            return TableData(
                table_id=f"tbl_{counter:03d}",
                caption=None,
                headers=headers,
                rows=data_rows,
            )
        except Exception:
            return None


# ------------------------------------------------------------------ #
#  Utilities                                                           #
# ------------------------------------------------------------------ #

# Map Word style names → heading level (1-6)
_HEADING_STYLES: dict[str, int] = {
    "Title": 1,
    "Subtitle": 2,
    **{f"Heading {i}": i for i in range(1, 7)},
}


def _heading_level(style_name: str) -> Optional[int]:
    """Return heading level (1-6) for heading styles, else None."""
    if style_name in _HEADING_STYLES:
        return _HEADING_STYLES[style_name]
    # Handle localised/stripped names like "heading 1" or "Heading1"
    m = re.match(r"^heading\s*(\d)$", style_name.strip().lower())
    if m:
        return int(m.group(1))
    return None


def _is_paragraph(block: object) -> bool:
    return type(block).__name__ == "Paragraph"


def _is_table(block: object) -> bool:
    return type(block).__name__ == "Table"


def _is_pseudo_heading(para: object) -> bool:
    """Detect pseudo-headings in documents that don't use Word Heading styles.

    A paragraph is treated as a pseudo-heading if ALL of:
      - Non-empty text
      - Short (≤ 100 characters)
    AND at least one of:
      - Text is entirely UPPER CASE (with at least one letter)
      - More than 50 % of characters are in bold runs
      - Text ends with ':' and contains no sentence-ending punctuation before it
        (i.e. it's a field label like "DETAILS OF INSURED PERSON:")

    This handles claim forms, policy documents, and other structured docs
    that use ALL CAPS Normal text as visual headings instead of Heading styles.
    """
    try:
        text: str = para.text.strip()  # type: ignore[attr-defined]
    except Exception:
        return False

    if not text or len(text) > 100:
        return False

    # ALL CAPS with at least one letter
    if text == text.upper() and any(c.isalpha() for c in text):
        return True

    # Majority of text in bold runs
    try:
        runs = para.runs  # type: ignore[attr-defined]
        total = sum(len(r.text) for r in runs if r.text)
        bold = sum(len(r.text) for r in runs if r.bold and r.text)
        if total > 0 and bold / total > 0.5:
            return True
    except Exception:
        pass

    # Field label: ends with ':' and no sentence-ending punctuation before it
    if (
        text.endswith(":")
        and "." not in text[:-1]
        and "?" not in text
        and "!" not in text
    ):
        return True

    return False


def _filename_to_title(stem: str) -> str:
    """Convert a filename stem to a readable title.

    Examples:
        project_brief_v2 → Project Brief V2
        meeting-notes    → Meeting Notes
    """
    return re.sub(r"[-_]+", " ", stem).title()
