"""Complex Tables · Step 4 — DOCX merged cells (test-first).

python-docx already returns a rectangular grid where merged cells repeat their value
(gridSpan across columns, vMerge down rows). The old parser DEDUPED consecutive equal
values, which misaligned merged columns AND collapsed legitimate duplicate values.
These tests pin the corrected behaviour. Skips if python-docx is not installed.

Run: pytest tests/test_docx_tables.py -v
"""
from __future__ import annotations

import pytest

docx = pytest.importorskip("docx", reason="python-docx not installed")
from docx import Document  # noqa: E402

from docnest.parsers.docx import DocxParser  # noqa: E402


class TestMergedCells:
    def test_horizontal_merge_repeats_value_aligned(self):
        d = Document()
        t = d.add_table(rows=2, cols=3)
        t.cell(0, 1).merge(t.cell(0, 2))            # header "Sales" spans cols 1-2
        t.cell(0, 0).text = "Region"
        t.cell(0, 1).text = "Sales"
        t.cell(1, 0).text = "Europe"; t.cell(1, 1).text = "10"; t.cell(1, 2).text = "20"
        td = DocxParser()._extract_table(t, 1)
        assert td is not None
        assert td.headers == ["Region", "Sales", "Sales"]      # not collapsed to 2 cols
        assert td.rows == [["Europe", "10", "20"]]
        assert all(len(r) == len(td.headers) for r in td.rows)

    def test_legitimate_duplicate_values_not_collapsed(self):
        d = Document()
        t = d.add_table(rows=2, cols=3)
        t.cell(0, 0).text = "Name"; t.cell(0, 1).text = "Q1"; t.cell(0, 2).text = "Q2"
        t.cell(1, 0).text = "Acme"; t.cell(1, 1).text = "10"; t.cell(1, 2).text = "10"
        td = DocxParser()._extract_table(t, 1)
        assert td.rows == [["Acme", "10", "10"]]               # both 10s preserved

    def test_vertical_merge_carries_value_down(self):
        d = Document()
        t = d.add_table(rows=3, cols=2)
        t.cell(1, 0).merge(t.cell(2, 0))            # vMerge col 0 across rows 1-2
        t.cell(0, 0).text = "Group"; t.cell(0, 1).text = "Item"
        t.cell(1, 0).text = "A"; t.cell(1, 1).text = "x"
        t.cell(2, 1).text = "y"
        td = DocxParser()._extract_table(t, 1)
        assert td.headers == ["Group", "Item"]
        # 'A' is carried into both rows it spans
        assert td.rows == [["A", "x"], ["A", "y"]]


class TestPlainTable:
    def test_plain_table_unaffected(self):
        d = Document()
        t = d.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "A"; t.cell(0, 1).text = "B"
        t.cell(1, 0).text = "1"; t.cell(1, 1).text = "2"
        td = DocxParser()._extract_table(t, 1)
        assert td.headers == ["A", "B"]
        assert td.rows == [["1", "2"]]
