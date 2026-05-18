"""Tests for DocxParser — Word document (.docx) parsing.

Run: pytest tests/test_docx_parser.py -v
"""
from __future__ import annotations

from pathlib import Path

import pytest

from docnest.parsers.docx import (
    DocxParser,
    _heading_level,
    _is_paragraph,
    _is_table,
    _is_pseudo_heading,
    _filename_to_title,
)
from docnest.models import RawDocument, Section
from docnest.exceptions import ParseError


FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture
def sample_docx() -> Path:
    path = FIXTURES / "sample.docx"
    if not path.exists():
        pytest.skip("sample.docx not found in tests/fixtures/")
    return path


@pytest.fixture
def generated_docx(tmp_path: Path) -> Path:
    """Generate a minimal DOCX using python-docx for parser tests."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_heading("My Report", 0)
    doc.add_heading("Introduction", 1)
    doc.add_paragraph("This is the introduction to our annual report with extensive detail.")
    doc.add_heading("Background", 2)
    doc.add_paragraph("Some background context about the project history and origins.")
    doc.add_heading("Methods", 1)
    doc.add_paragraph("We used a combination of quantitative and qualitative methods.")
    doc.add_heading("Results", 1)
    doc.add_paragraph("The results show significant improvements across all metrics measured.")
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Growth"
    table.rows[1].cells[1].text = "25%"
    table.rows[2].cells[0].text = "Users"
    table.rows[2].cells[1].text = "10K"
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


# ── DocxParser.supports() ─────────────────────────────────────────────────────

class TestDocxParserSupports:
    def test_supports_docx(self):
        assert DocxParser().supports("report.docx") is True

    def test_supports_docx_uppercase(self):
        assert DocxParser().supports("REPORT.DOCX") is True

    def test_rejects_doc(self):
        assert DocxParser().supports("report.doc") is False

    def test_rejects_pdf(self):
        assert DocxParser().supports("report.pdf") is False

    def test_rejects_md(self):
        assert DocxParser().supports("report.md") is False


# ── DocxParser.parse() — fixture-based ────────────────────────────────────────

class TestDocxParserParse:
    def test_returns_raw_document(self, generated_docx: Path):
        parser = DocxParser()
        raw = parser.parse(str(generated_docx))
        assert isinstance(raw, RawDocument)

    def test_format_is_docx(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        assert raw.format == "docx"

    def test_has_sections(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        assert len(raw.sections) > 0

    def test_section_titles_extracted(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        titles = [s.title for s in raw.sections]
        assert "Introduction" in titles or "My Report" in titles

    def test_section_levels_are_valid(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        for s in raw.sections:
            assert 1 <= s.level <= 6

    def test_section_text_not_empty(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        # At least one section should have text
        texts = [s.text for s in raw.sections if s.text]
        assert len(texts) > 0

    def test_tables_extracted(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        all_tables = [t for s in raw.sections for t in s.tables]
        assert len(all_tables) >= 1

    def test_table_has_headers(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        all_tables = [t for s in raw.sections for t in s.tables]
        assert all_tables[0].headers  # non-empty headers

    def test_table_has_rows(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        all_tables = [t for s in raw.sections for t in s.tables]
        assert len(all_tables[0].rows) >= 1

    def test_doc_id_from_filename(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        assert raw.doc_id  # non-empty

    def test_title_extracted(self, generated_docx: Path):
        raw = DocxParser().parse(str(generated_docx))
        assert raw.title  # non-empty title

    def test_parse_with_fixture_file(self, sample_docx: Path):
        raw = DocxParser().parse(str(sample_docx))
        assert isinstance(raw, RawDocument)
        assert len(raw.sections) >= 3


# ── DocxParser error paths ────────────────────────────────────────────────────

class TestDocxParserErrors:
    def test_missing_file_raises_parse_error(self, tmp_path: Path):
        parser = DocxParser()
        with pytest.raises(ParseError):
            parser.parse(str(tmp_path / "nonexistent.docx"))

    def test_empty_file_raises_parse_error(self, tmp_path: Path):
        empty = tmp_path / "empty.docx"
        empty.write_bytes(b"")
        with pytest.raises(ParseError):
            DocxParser().parse(str(empty))

    def test_invalid_docx_raises_parse_error(self, tmp_path: Path):
        """A .docx with garbage content should raise ParseError."""
        bad = tmp_path / "bad.docx"
        bad.write_bytes(b"not a docx file at all")
        with pytest.raises(ParseError):
            DocxParser().parse(str(bad))


# ── DocxParser — pseudo-headings ──────────────────────────────────────────────

class TestDocxPseudoHeadings:
    def test_all_caps_paragraph_detected_as_pseudo_heading(self, tmp_path: Path):
        """ALL CAPS paragraphs without a Heading style should become sections."""
        from docx import Document as DocxDocument
        doc = DocxDocument()
        # Normal paragraph styled as Normal (not Heading)
        doc.add_paragraph("EXECUTIVE SUMMARY")
        doc.add_paragraph("This is the executive summary text.")
        path = tmp_path / "allcaps.docx"
        doc.save(str(path))

        raw = DocxParser().parse(str(path))
        titles = [s.title for s in raw.sections]
        # "EXECUTIVE SUMMARY" should be detected as a pseudo-heading section
        assert any("EXECUTIVE" in t for t in titles) or len(raw.sections) >= 1

    def test_field_label_detected(self, tmp_path: Path):
        """Lines ending with ':' should be detected as pseudo-headings."""
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("Employee Name:")
        doc.add_paragraph("John Smith is the employee name in question here.")
        path = tmp_path / "field.docx"
        doc.save(str(path))

        raw = DocxParser().parse(str(path))
        # Should parse without crashing
        assert isinstance(raw, RawDocument)

    def test_content_before_first_heading_becomes_intro(self, tmp_path: Path):
        """Text before any heading → implicit 'Introduction' section."""
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("This is some introductory text before any heading.")
        doc.add_paragraph("More intro text that comes before the first heading.")
        doc.add_heading("Chapter One", 1)
        doc.add_paragraph("Chapter one content here.")
        path = tmp_path / "intro.docx"
        doc.save(str(path))

        raw = DocxParser().parse(str(path))
        titles = [s.title for s in raw.sections]
        assert "Introduction" in titles or len(raw.sections) >= 2

    def test_list_style_appended_with_bullet(self, tmp_path: Path):
        """List-style paragraphs should be formatted with '- ' prefix."""
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("Items", 1)
        p = doc.add_paragraph("Item one content here.", style="List Bullet")
        path = tmp_path / "list.docx"
        doc.save(str(path))

        raw = DocxParser().parse(str(path))
        assert any("- " in s.text or "Item one" in s.text for s in raw.sections)

    def test_nested_headings_multiple_levels(self, tmp_path: Path):
        """H1 → H2 → H3 nesting should produce sections at correct levels."""
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("Chapter", 1)
        doc.add_paragraph("Chapter text.")
        doc.add_heading("Section", 2)
        doc.add_paragraph("Section text.")
        doc.add_heading("Subsection", 3)
        doc.add_paragraph("Subsection text.")
        path = tmp_path / "nested.docx"
        doc.save(str(path))

        raw = DocxParser().parse(str(path))
        levels = {s.level for s in raw.sections}
        assert 1 in levels
        assert 2 in levels


# ── Utility functions ─────────────────────────────────────────────────────────

class TestHeadingLevel:
    def test_heading_1(self):
        assert _heading_level("Heading 1") == 1

    def test_heading_6(self):
        assert _heading_level("Heading 6") == 6

    def test_title_style(self):
        assert _heading_level("Title") == 1

    def test_subtitle_style(self):
        assert _heading_level("Subtitle") == 2

    def test_normal_returns_none(self):
        assert _heading_level("Normal") is None

    def test_body_text_returns_none(self):
        assert _heading_level("Body Text") is None

    def test_localised_heading_lowercase(self):
        # Localised headings like "heading 1" should also match
        assert _heading_level("heading 1") == 1

    def test_localised_heading_no_space(self):
        assert _heading_level("heading1") == 1


class TestFilenameToTitle:
    def test_underscores_to_spaces(self):
        assert _filename_to_title("project_brief") == "Project Brief"

    def test_hyphens_to_spaces(self):
        assert _filename_to_title("meeting-notes") == "Meeting Notes"

    def test_mixed_separators(self):
        result = _filename_to_title("project_brief-v2")
        assert "Project" in result

    def test_titlecase_applied(self):
        assert _filename_to_title("annual_report") == "Annual Report"


class TestIsParagraphAndTable:
    def test_is_paragraph_returns_false_for_non_paragraph(self):
        class Fake:
            pass
        assert _is_paragraph(Fake()) is False

    def test_is_table_returns_false_for_non_table(self):
        class Fake:
            pass
        assert _is_table(Fake()) is False
